import io
import docx 

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import resolve1
from requests.exceptions import ConnectionError

from googletrans import Translator


def pdf_to_text(file, pages=None):
    if not pages:
        pagenums = set()
    else:
        pagenums = set(pages)

    output = io.StringIO()
    manager = PDFResourceManager()
    converter = TextConverter(manager, output, laparams=LAParams())
    interpreter = PDFPageInterpreter(manager, converter)

    infile = open(file, 'rb')
    for page in PDFPage.get_pages(infile, pagenums):
        interpreter.process_page(page)
    infile.close()
    converter.close()
    text = output.getvalue()
    output.close
    return text 


def getText(file):
    doc = docx.Document(file)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)


def get_pdf_totalpage(file):
    file = open(file, 'rb')
    parser = PDFParser(file)
    document = PDFDocument(parser)
    page_count = resolve1(document.catalog['Pages'])['Count']
    return page_count


class ExtractMixin(object):


    def extract_pdf(self, learning):
        pdf_content = []
        page_count = get_pdf_totalpage(learning.original_file.path)

        for number_of_pages in range(page_count):
            content = pdf_to_text(learning.original_file.path, pages=[number_of_pages,page_count])
            pdf_content.append(content)

        original_content =' '.join(pdf_content)

        return original_content


    def extract_docx(self, learning):
        content = getText(learning.original_file.path)

        original_content = content
        return original_content


class DetectLanguageMixin(object):
    translator = Translator()


    def detect_lang(self, content):
        try: 
            extracted_content = [] 
            content_count = len(content)
            min = 200
            for min in range(content_count):
                source_lang = self.translator.detect(content[1:min])
                return source_lang.lang
        except ConnectionError as e:
            return 'ConnectionError'


class TranslationMixin(object):
    translator = Translator()


    def translate(self, content):
        try:
            content_count = len(content)
            translation = []
            if content_count >= 5000:
                start = 0
                end = 700
                for end in range(content_count):
                    translated = self.translator.translate(content[start:end],dest='ceb')

                    translation.append(translated.text)
                    start += end    
                    end += 700
                translated_content = ' '.join(translation)
                return translated_content
            elif content_count == 0:
                return 'File is Empty.'
            else :
                translated_content = self.translator.translate(content,dest='ceb')
                return translated_content.text
        except ConnectionError as e:
            return 'ConnectionError'